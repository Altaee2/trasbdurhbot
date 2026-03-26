import os
import json
import telebot
from telebot import types
from deep_translator import GoogleTranslator
from docx import Document
import fitz  # PyMuPDF
from fpdf import FPDF
from arabic_reshaper import reshape
from bidi.algorithm import get_display

# --- الإعدادات الأساسية ---
API_TOKEN = '8736450911:AAHqY6q6W8EBipZFtzq2OrYUvn-CKwctEkA'
ADMIN_ID = 6454550864  # ايديك هنا
bot = telebot.TeleBot(API_TOKEN)
bot_namee = "@ALAAHI12BOT"
DB_FILE = "bot_data.json"

def load_data():
    if not os.path.exists(DB_FILE):
        return {"users": [], "ban": [], "channels": {}, "sub_msg": "⚠️ يجب عليك الاشتراك في قنوات البوت أولاً لتتمكن من استخدامه!"}
    with open(DB_FILE, "r") as f:
        return json.load(f)

def save_data(data):
    with open(DB_FILE, "w") as f:
        json.dump(data, f, indent=4)

data = load_data()
user_files_temp = {}

# --- وظائف التحقق والاشتراك ---
def is_sub(user_id):
    if not data["channels"]: return True
    for ch_id in data["channels"]:
        try:
            status = bot.get_chat_member(ch_id, user_id).status
            if status in ["left", "kicked"]: return False
        except: continue
    return True

def format_arabic_for_pdf(text):
    try:
        reshaped_text = reshape(text)
        return get_display(reshaped_text)
    except: return text

def split_text(text, chunk_size=3000):
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

# --- لوحات المفاتيح ---
def admin_keyboard():
    markup = types.InlineKeyboardMarkup(row_width=2)
    markup.add(
        types.InlineKeyboardButton("📊 الإحصائيات", callback_data="adm_stats"),
        types.InlineKeyboardButton("📢 الإذاعة", callback_data="adm_bc"),
        types.InlineKeyboardButton("🔒 الاشتراك الإجباري", callback_data="adm_ch"),
        types.InlineKeyboardButton("🚫 حظر/إلغاء حظر", callback_data="adm_ban"),
        types.InlineKeyboardButton("📝 رسالة الاشتراك", callback_data="adm_msg")
    )
    return markup
def escape_html(text):
    if not text: return "None"
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

# --- الأوامر الأساسية ---
import html

def escape_html(text):
    """دالة لتنظيف النص من رموز HTML لمنع توقف الإشعار"""
    return html.escape(text)
def create_rights_keyboard():
    # جلب معلومات البوت للحصول على اليوزر واسم البوت تلقائياً
    bot_info = bot.get_me()
    bot_username = bot_info.username
    bot_name = bot_info.first_name

    keyboard = types.InlineKeyboardMarkup(row_width=1)
    
    # الزر الأول (الرابط الخاص بك)
    btn_vip = types.InlineKeyboardButton("✔️ خدمات احترافية VIP لطلبة الجامعات 👇", url="https://t.me/+9u0H97uCRT5hM2Vi")
    
    # نص المشاركة الذي سيظهر عند إرساله لصديق
    share_text = f"جرب هذا البوت الرهيب {bot_name} 🤖\nيقوم بترجمة ملفات PDF و DOCX بدقة عالية ويحافظ على التنسيق! ✨\n\nرابط البوت: t.me/{bot_username}"
    
    # زر المشاركة الحقيقي
    # استخدام switch_inline_query_current_chat يجعل المستخدم يختار دردشة لإرسال النص إليها
    btn_share = types.InlineKeyboardButton("📤 مشاركة البوت مع أصدقائك", switch_inline_query=share_text)
    
    keyboard.add(btn_vip, btn_share)
    return keyboard

@bot.message_handler(commands=['start'])
def start(message):
    uid = message.from_user.id
    # تحميل البيانات للتأكد من تحديثها
    global data
    data = load_data()

    # 1. فحص الحظر
    if uid in data.get("ban", []):
        bot.send_message(message.chat.id, "❌ تم حظرك من استخدام البوت.")
        return
    if not is_sub(uid):
        markup = types.InlineKeyboardMarkup(row_width=1)
        if data["channels"]:
            for ch_id, ch_link in data["channels"].items():
                try:
                    # جلب معلومات القناة للحصول على اسمها الحقيقي
                    ch_info = bot.get_chat(ch_id)
                    ch_title = ch_info.title
                    button_text = f"{ch_title}"
                except:
                    # في حال فشل الجلب (مثلاً البوت ليس أدمن) نضع نص احتياطي
                    button_text = "قناة الاشتراك 📢"
                
                markup.add(types.InlineKeyboardButton(button_text, url=ch_link))
        
        markup.add(types.InlineKeyboardButton("تم الاشتراك ✅", callback_data="check_sub"))
        
        # إرسال الرسالة مع الأزرار التي تحمل أسماء القنوات
        bot.send_message(uid, data["sub_msg"], reply_markup=markup)
        return

    # 2. فحص المستخدم الجديد وإرسال الإشعار
    if uid not in data.get("users", []):
        data["users"].append(uid)
        save_data(data)
        
        # تنظيف البيانات للإشعار
        first_name = message.from_user.first_name if message.from_user.first_name else ""
        last_name = message.from_user.last_name if message.from_user.last_name else ""
        name = escape_html(f"{first_name} {last_name}".strip())
        username = f"@{message.from_user.username}" if message.from_user.username else "لا يوجد"
        log_text = (f"🚀 <b>مستخدم جديد دخل للبوت!</b>\n\n"
                    f"👤 الاسم: <b>{name}</b>\n"
                    f"🔗 اليوزر: <b>{username}</b>\n"
                    f"🆔 الآيدي: <code>{uid}</code>\n")
        try:
            bot.send_message(ADMIN_ID, log_text, parse_mode="HTML")
        except:
            bot.send_message(ADMIN_ID, f"🚀 مستخدم جديد!\nالاسم: {name}\nالايدي: {uid}\n اليوزر: {username}\n")
        
    if uid == ADMIN_ID:
        bot.send_message(uid, "💎 أهلاً بك في لوحة تحكم المطور:", reply_markup=admin_keyboard())
        

    
    bot.reply_to(message, f"أهلاً بك {message.from_user.first_name}! أرسل لي ملف PDF أو DOCX وسأقوم بترجمته.")
@bot.message_handler(commands=['admin'])
def admin_panel(message):
    if message.from_user.id == ADMIN_ID:
        bot.send_message(message.chat.id, "💎 أهلاً بك في لوحة تحكم المطور:", reply_markup=admin_keyboard())

def broadcast_step(message):
    count = 0
    for u in data["users"]:
        try:
            bot.send_message(u, message.text)
            count += 1
        except: continue
    bot.send_message(ADMIN_ID, f"✅ تم إرسال الإذاعة لـ {count} مستخدم.")

def add_channel_step(message):
    input_data = message.text.strip()
    
    # 1. تنظيف المدخلات (استخراج اليوزر من الرابط إذا وجد)
    target = input_data
    if "t.me/" in input_data:
        target = input_data.split("t.me/")[-1].replace("@", "")
    elif not input_data.startswith("-100") and not input_data.startswith("@"):
        target = f"@{input_data}"

    try:
        # 2. محاولة جلب معلومات الدردشة (تحويل اليوزر/الرابط إلى ID)
        target_chat = bot.get_chat(target)
        ch_id = str(target_chat.id)
        
        # 3. تجهيز الرابط (إذا كان يوزر نسوي رابط، إذا كان رابط خاص ننزله نفسه)
        ch_link = f"https://t.me/{target_chat.username}" if target_chat.username else input_data
        
        # 4. حفظ البيانات
        data["channels"][ch_id] = ch_link
        save_data(data)
        
        bot.send_message(
            ADMIN_ID, 
            f"✅ **تمت إضافة القناة بنجاح!**\n\n"
            f"📌 الاسم: {target_chat.title}\n"
            f"🆔 الايدي: `{ch_id}`\n"
            f"🔗 الرابط: {ch_link}", 
            parse_mode="Markdown"
        )
    except Exception as e:
        # رسالة خطأ واضحة في حال فشل البوت في الوصول للقناة
        error_msg = (f"❌ **فشل إضافة القناة!**\n\n"
                     f"السبب المحتمل:\n"
                     f"1. البوت ليس أدمن في القناة (ضروري جداً).\n"
                     f"2. اليوزر أو الرابط خطأ.\n\n"
                     f"⚠️ الخطأ التقني: `{e}`")
        bot.send_message(ADMIN_ID, error_msg, parse_mode="Markdown")
def ban_step(message):
    try:
        target = int(message.text)
        if target in data["ban"]:
            data["ban"].remove(target); msg = "✅ تم إلغاء الحظر."
        else:
            data["ban"].append(target); msg = "🚫 تم الحظر."
        save_data(data); bot.send_message(ADMIN_ID, msg)
    except: bot.send_message(ADMIN_ID, "⚠️ ارسل ايدي صحيح.")

def change_msg_step(message):
    data["sub_msg"] = message.text
    save_data(data)
    bot.send_message(ADMIN_ID, "✅ تم تحديث رسالة الاشتراك.")

# --- Callback Handler ---
@bot.callback_query_handler(func=lambda call: True)
def callback_all(call):
    uid = call.message.chat.id
    
    if call.data == "check_sub":
        if is_sub(uid):
            bot.delete_message(uid, call.message.message_id)
            bot.send_message(uid, "✅ تم التحقق، يمكنك استخدام البوت الآن.\n ارسل /start من جديد ")
        else:
            bot.answer_callback_query(call.id, "❌ لم تشترك بعد!", show_alert=True)

    elif uid == ADMIN_ID:
        if call.data == "adm_stats":
            msg = f"📊 **إحصائيات البوت:**\n\n👥 المستخدمين: {len(data['users'])}\n🚫 المحظورين: {len(data['ban'])}\n📢 القنوات: {len(data['channels'])}"
            bot.send_message(uid, msg, parse_mode="Markdown")
        
        elif call.data == "adm_bc":
            bot.register_next_step_handler(bot.send_message(uid, "ارسل رسالة الإذاعة:"), broadcast_step)
        
        elif call.data == "adm_ch":
            msg = "📢 **قنوات الاشتراك الإجباري الحالية:**\n\n"
            for ch_id, ch_link in data["channels"].items():
                try:
                    # جلب معلومات القناة حياً من تلجرام
                    ch_info = bot.get_chat(ch_id)
                    ch_name = ch_info.title
                except:
                    ch_name = "قناة غير معروفة (تأكد من وجود البوت كأدمن)"

                msg += (f"📌 **الاسم:** {ch_name}\n"
                        f"🆔 **الايدي:** `{ch_id}`\n"
                        f"🔗 **الرابط:** [اضغط هنا]({ch_link})\n"
                        f"─────────────────\n")
            kb = types.InlineKeyboardMarkup()
            kb.add(types.InlineKeyboardButton("➕ إضافة قناة", callback_data="add_ch"),
                   types.InlineKeyboardButton("🗑 مسح الكل", callback_data="clear_ch"))
            bot.send_message(uid, msg, reply_markup=kb, parse_mode="Markdown")
        
        elif call.data == "add_ch":
            bot.register_next_step_handler(bot.send_message(uid, "ارسل يوزر القناة المراد اضافتها للاشتراك الاجباري "), add_channel_step)
        
        elif call.data == "clear_ch":
            data["channels"] = {}; save_data(data)
            bot.answer_callback_query(call.id, "✅ تم حذف جميع القنوات.")
            bot.edit_message_text("✅ تم مسح جميع قنوات الاشتراك الإجباري.", uid, call.message.message_id)

        elif call.data == "adm_ban":
            bot.register_next_step_handler(bot.send_message(uid, "ارسل ايدي المستخدم:"), ban_step)
        
        elif call.data == "adm_msg":
            bot.register_next_step_handler(bot.send_message(uid, "ارسل النص الجديد لرسالة الاشتراك الاجباري:"), change_msg_step)

    if call.data.startswith("lang_"):
        process_translation(call, call.data.split("_")[1])

# --- معالجة الملفات والترجمة ---
@bot.message_handler(content_types=['document'])
def handle_docs(message):
    uid = message.chat.id
    if uid in data["ban"] or not is_sub(uid): return start(message)
    
    fn = message.document.file_name
    ext = os.path.splitext(fn)[1].lower()
    if ext not in ['.pdf', '.docx']: return bot.reply_to(message, "⚠️ غير مدعوم فقط PDF AND DOCX.")

    m = bot.reply_to(message, "📥 جاري الاستلام... انتظر لأختيارك للغة")
    fi = bot.get_file(message.document.file_id)
    dl = bot.download_file(fi.file_path)
    with open(fn, 'wb') as f: f.write(dl)

    p = 0
    if ext == '.pdf':
        with fitz.open(fn) as d: p = len(d)
    else:
        d = Document(fn); p = len(d.paragraphs)

    user_files_temp[uid] = {'file': fn, 'size': f"{message.document.file_size/1024:.2f} KB", 'pages': p}
    bot.delete_message(uid, m.message_id)
    
    kb = types.InlineKeyboardMarkup()
    kb.add(types.InlineKeyboardButton("العربية 🇸🇦", callback_data='lang_ar'),
           types.InlineKeyboardButton("English 🇺🇸", callback_data='lang_en'))
    bot.send_message(uid, "✅ اختر لغة الترجمة:", reply_markup=kb)

def process_translation(call, lang):
    uid = call.message.chat.id
    if uid not in user_files_temp: return
    
    info = user_files_temp[uid]
    in_f = info['file']
    out_f = f"translated_{uid}.pdf"
    
    bot.edit_message_text("⏳ جاري الترجمة والتحويل انتظر قليلا من فضلك..", uid, call.message.message_id)
    
    try:
        txt = ""
        if in_f.endswith('.pdf'):
            with fitz.open(in_f) as d:
                for pg in d: txt += pg.get_text()
        else:
            doc = Document(in_f)
            txt = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        chunks = split_text(txt)
        # ترجمة المقاطع
        tr = [GoogleTranslator(source='auto', target=lang).translate(c) for c in chunks if c.strip()]
        final = "\n".join(tr)

        # إنشاء ملف PDF
        pdf = FPDF()
        pdf.add_page()
        if os.path.exists("arial.ttf"):
            pdf.add_font("CustomFont", "", "arial.ttf", uni=True) # أضف uni=True هنا
            pdf.set_font("CustomFont", size=12)
        else:
            pdf.set_font("Arial", size=12)

        if lang == 'ar':
            pdf.multi_cell(0, 10, txt=format_arabic_for_pdf(final), align='R')
        else:
            pdf.multi_cell(0, 10, txt=final, align='L')

        pdf.output(out_f)

        # تحضير الكابشن المطلوب
        caption = (f"✅ تم إكمال الترجمة بنجاح\n\n"
                   f"📄 الملف: {in_f}\n"
                   f"🌐 اللغة: {lang}\n\n"
                   f"تم الترجمة بواسطة بوت : {bot_namee}\n")

        # إرسال الملف مع الكابشن ولوحة الحقوق
        with open(out_f, 'rb') as f:
            bot.send_document(uid, f, caption=caption, reply_markup=create_rights_keyboard())
            
    except Exception as e:
        bot.send_message(uid, f"❌ فشلت عملية الترجمة: {str(e)}")
    
    # تنظيف الملفات المؤقتة
    if os.path.exists(in_f): os.remove(in_f)
    if os.path.exists(out_f): os.remove(out_f)
    if uid in user_files_temp: del user_files_temp[uid]
print(" البوت والأدمن يعملان بشكل صحيح...")
bot.infinity_polling()
